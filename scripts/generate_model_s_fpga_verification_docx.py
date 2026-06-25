from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


RESULTS_DIR = Path(
    r"C:\Users\YangGeon\Desktop\작업물\학업\한양대학교 문서\스펙\대회 관련\2026 전국 반도체 설계대전\Results"
)
OUT_DOCX = RESULTS_DIR / "Model_S_FPGA_Verification_Report.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "E8EEF5"
GREEN = "E2F0D9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(hdr.cells[i], LIGHT_GRAY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            if i > 0 and str(value).replace(".", "").replace("%", "").replace("/", "").replace(" ", "").replace("-", "").isdigit():
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            run.font.size = Pt(9)
    if widths:
        set_table_width(table, widths)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_callout(doc, label, text, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_table_width(table, [9360])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + " ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph()


def configure_styles(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def main():
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Model S FPGA 검증 보고서")
    r.bold = True
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    s = subtitle.add_run("SNN-inspired ECG 4-class RTL Classifier / Nexys A7-100T 구현 검증")
    s.font.size = Pt(11)
    s.font.color.rgb = RGBColor(80, 80, 80)

    add_table(
        doc,
        ["항목", "내용"],
        [
            ["검증 대상", "Model S RTL core 및 Nexys A7 보드 데모 bitstream"],
            ["대상 FPGA", "Xilinx Artix-7 xc7a100tcsg324-1, Nexys A7-100T"],
            ["Vivado", "Vivado 2020.2"],
            ["통합 프로젝트", r"C:\Users\YangGeon\SNN_ECG_RESTORE_MODEL_S\vivado_project\SNN_ECG_ModelS_Unified\SNN_ECG_ModelS_Unified.xpr"],
            ["보드 top", "nexys_a7_model_s_smoke_top"],
            ["Core top", "snn_ecg_model_a_plus_core"],
            ["Bitstream", r"C:\Users\YangGeon\SNN_ECG_RESTORE_MODEL_S\bitstreams\nexys_a7_model_s_smoke_top.bit"],
        ],
        widths=[2100, 7260],
    )

    add_callout(
        doc,
        "검증 결론.",
        "Model S RTL은 Vivado에서 합성 및 구현을 통과했고, Nexys A7-100T 보드에 bitstream으로 program되어 버튼/7-segment 기반 interactive demo가 동작했다. Timing violation과 DRC violation은 없으며, Model S core 기준 DSP와 BRAM 사용량은 0이다.",
        GREEN,
    )

    add_heading(doc, "1. 검증 범위", 1)
    add_body(
        doc,
        "본 문서는 Model S ECG classifier RTL이 FPGA에서 구현 가능한지 확인한 결과를 정리한다. 검증은 두 층위로 나누어 해석한다. 첫째, 실제 classifier logic인 Model S core의 자원 사용량을 확인한다. 둘째, Nexys A7 보드에서 버튼과 7-segment display를 사용해 동작을 확인하기 위한 board smoke wrapper의 구현 결과를 확인한다.",
    )
    add_body(
        doc,
        "자원 사용량은 classifier의 실제 SoC/IP 규모를 판단해야 하므로 wrapper를 제외한 core 기준 값을 우선 지표로 사용한다. 반면 전력, 온도, timing, DRC는 실제 bitstream으로 구현된 board demo 전체 기준 보고서를 함께 제시한다.",
    )

    add_heading(doc, "2. Model S Core 자원 사용량", 1)
    add_body(
        doc,
        "아래 표는 `snn_ecg_model_a_plus_core`를 top으로 합성한 결과다. 이 값은 7-segment, button controller, demo ECG ROM을 제외한 classifier core 기준이다. 따라서 실제 classifier IP의 크기를 판단할 때는 이 표를 기준으로 보는 것이 맞다.",
    )
    add_table(
        doc,
        ["Resource", "Used", "Available", "Utilization", "해석"],
        [
            ["Slice LUTs", "5309", "63400", "8.37%", "pNN/RDM/DSCR/RAM/ECP/RBBB/EERG logic 포함"],
            ["Slice Registers", "1250", "126800", "0.99%", "상태 register 및 feature counter 포함"],
            ["Block RAM Tile", "0", "135", "0.00%", "core 내부는 BRAM-free"],
            ["DSPs", "0", "240", "0.00%", "곱셈기/DSP-free 구조 유지"],
            ["Bonded IOB", "21", "210", "10.00%", "core 단독 입출력 포트 기준"],
            ["BUFGCTRL", "1", "32", "3.13%", "core clock buffer"],
        ],
        widths=[1900, 1100, 1300, 1300, 3760],
    )
    add_callout(
        doc,
        "핵심 해석.",
        "Model S core는 LUT 8.37%, FF 0.99%, BRAM 0%, DSP 0%로 합성된다. 즉 ECG classifier 자체는 multiplier-free, BRAM-free, 비교기/카운터/레지스터 중심의 저자원 RTL 구조로 구현되어 있다.",
    )

    add_heading(doc, "3. Board Smoke Bitstream 자원 사용량", 1)
    add_body(
        doc,
        "보드 검증 bitstream은 core 외에 버튼 입력, 7-segment display controller, class별 60초 ECG 예시 ROM을 포함한다. 따라서 board smoke bitstream의 BRAM 사용량은 classifier core의 BRAM 사용량이 아니라 demo dataset ROM 비용이다.",
    )
    add_table(
        doc,
        ["Resource", "Used", "Available", "Utilization", "비고"],
        [
            ["Slice LUTs", "5400", "63400", "8.52%", "core + board wrapper"],
            ["Slice Registers", "1384", "126800", "1.09%", "core + display/control state"],
            ["Block RAM Tile", "84", "135", "62.22%", "4개 60초 ECG 예시 ROM 때문"],
            ["DSPs", "0", "240", "0.00%", "보드 wrapper 포함해도 DSP 사용 없음"],
            ["Bonded IOB", "39", "210", "18.57%", "버튼, LED, 7-segment, clock/reset 포함"],
            ["BUFGCTRL", "2", "32", "6.25%", "100 MHz board clock 및 1 MHz core clock"],
        ],
        widths=[1900, 1100, 1300, 1300, 3760],
    )
    add_body(
        doc,
        "BRAM 84개는 FPGA 안에 예시 ECG `.mem` 데이터를 저장하기 위한 비용이다. 실제 ADC stream을 외부에서 입력받는 최종 칩 형태에서는 이 ROM이 제거되므로 classifier core 기준 BRAM 사용량은 0으로 보는 것이 타당하다.",
    )

    add_heading(doc, "4. Timing / DRC 검증", 1)
    add_table(
        doc,
        ["항목", "결과", "해석"],
        [
            ["DRC", "No DRC violations", "보드 핀 제약 및 I/O standard 문제가 해결됨"],
            ["Setup timing", "WNS +4.242 ns, TNS 0 ns", "timing violation 없음"],
            ["Hold timing", "WHS +0.085 ns, THS 0 ns", "hold violation 없음"],
            ["Failing endpoints", "0", "구현 후 timing 실패 endpoint 없음"],
            ["Clock 구조", "100 MHz board clock -> 1 MHz core clock", "ECG 1 kSPS 처리보다 충분히 빠른 내부 검증 clock"],
        ],
        widths=[2100, 2500, 4760],
    )
    add_body(
        doc,
        "Project Summary 기준으로 implementation은 완료 상태이며, timing constraints는 만족한다. 따라서 현재 bitstream은 Nexys A7-100T에 program 가능한 정상 구현물로 판단한다.",
    )

    add_heading(doc, "5. 전력 및 온도 추정", 1)
    add_table(
        doc,
        ["항목", "값", "해석"],
        [
            ["Total On-Chip Power", "0.104 W", "Vivado vectorless 추정치"],
            ["Dynamic Power", "0.004 W", "core가 1 MHz로 동작하고 toggle이 낮아 작게 추정됨"],
            ["Device Static Power", "0.101 W", "Artix-7 정적 전력 성분"],
            ["Junction Temperature", "25.5 C", "추정 온도, 열 여유 충분"],
            ["Thermal Margin", "59.5 C", "온도 한계까지 여유 있음"],
            ["Effective thetaJA", "4.6 C/W", "Vivado board/thermal model 기준"],
            ["Confidence Level", "Low", "입력 activity가 실제 측정이 아니므로 참고값"],
        ],
        widths=[2500, 1900, 4960],
    )
    add_body(
        doc,
        "전력 수치는 실제 전류 측정값이 아니라 Vivado의 activity 추정 기반 결과다. 특히 Confidence Level이 Low이므로 절대 전력값으로 주장하기보다는, thermal risk가 없고 구현 규모가 작다는 참고 지표로 해석한다.",
    )

    add_heading(doc, "6. 보드 기능 검증", 1)
    add_body(
        doc,
        "보드 데모는 외부 ADC 대신 FPGA 내부 ROM에 저장된 strict-test 60초 ECG 예시 4개를 사용한다. 버튼을 누르면 해당 class 예시가 Model S core로 stream되고, 예측 결과가 7-segment display에 표시된다.",
    )
    add_table(
        doc,
        ["버튼", "입력 예시", "표시 동작"],
        [
            ["BTNU", "NSR example", "왼쪽 4자리: 예측 class, 오른쪽 4자리: CORR/ERR"],
            ["BTNL", "ARR example", "새 버튼 입력 시 이전 결과는 blank 처리"],
            ["BTND", "CHF example", "결과가 나오기 전까지 7-segment 비활성"],
            ["BTNR", "AFF example", "pred_valid 이후 결과 latch"],
            ["BTNC", "pseudo-random class", "4개 class 중 하나를 pseudo-random 선택"],
        ],
        widths=[1300, 2200, 5860],
    )
    add_body(
        doc,
        "초기 구현에서는 ECG sample 사이에 idle clock이 들어가 CHF가 NSR로, ARR가 AFF로 오분류되는 문제가 있었다. XSim timing debug에서 sample gap이 원인임을 확인했고, board wrapper를 strict XSim dataset testbench와 동일하게 한 core clock마다 한 sample을 공급하도록 수정했다. 이 수정 후 보드 demo 동작은 simulation timing과 일치한다.",
    )

    add_heading(doc, "7. RTL 정확도 검증 결과", 1)
    add_body(
        doc,
        "분류 정확도는 Vivado Project Summary가 아니라 strict record-wise XSim dataset 검증 결과를 기준으로 한다. Model S RTL 내부에 EERG 및 RBBB-delay readout이 포함된 상태에서 train/validation/test split을 모두 XSim으로 검증했다.",
    )
    add_table(
        doc,
        ["Split", "Segment Accuracy", "Record Accuracy", "Macro-F1", "Balanced Accuracy"],
        [
            ["Train", "313/400 = 78.25%", "41/50 = 82.00%", "78.22%", "78.25%"],
            ["Validation", "136/160 = 85.00%", "18/20 = 90.00%", "84.91%", "85.00%"],
            ["Test", "131/160 = 81.88%", "18/19 = 94.74%", "81.93%", "81.88%"],
        ],
        widths=[1800, 2300, 2200, 1500, 1560],
    )
    add_table(
        doc,
        ["Actual \\ Pred", "NSR", "CHF", "ARR", "AFF"],
        [
            ["NSR", "31", "0", "9", "0"],
            ["CHF", "0", "37", "3", "0"],
            ["ARR", "6", "0", "28", "6"],
            ["AFF", "0", "3", "2", "35"],
        ],
        widths=[2200, 1790, 1790, 1790, 1790],
    )
    add_body(
        doc,
        "위 결과는 record-wise holdout test 기준이다. 따라서 같은 record에서 잘린 segment가 train/validation/test에 섞이는 leakage를 피한 상태에서 얻은 성능으로 해석한다.",
    )

    add_heading(doc, "8. 최종 판단", 1)
    add_bullets(
        doc,
        [
            "Model S core는 DSP 0%, BRAM 0%로 합성되어 SNN-inspired low-resource RTL 구조를 유지한다.",
            "Nexys A7 board demo bitstream은 DRC 및 timing을 통과했고 실제 보드 program까지 완료되었다.",
            "보드 bitstream의 BRAM 62.22%는 4개 ECG 예시 ROM 때문이며, classifier core의 고유 BRAM 비용이 아니다.",
            "Vivado 전력 추정은 0.104 W, junction temperature 25.5 C로 thermal risk는 낮지만, confidence level이 Low이므로 참고값으로만 사용한다.",
            "분류 성능은 strict record-wise XSim test 기준 segment 81.88%, record 94.74%로 보고한다.",
        ],
    )
    add_callout(
        doc,
        "최종 결론.",
        "현재 통합 프로젝트와 bitstream은 Model S RTL의 FPGA 구현 가능성, timing closure, board-level interactive operation을 확인한 검증 산출물이다. 실제 외부 ADC 연동 검증은 별도 단계로 남지만, RTL core 자체와 보드 demo bitstream은 FPGA 검증을 통과한 상태로 판단한다.",
        GREEN,
    )

    add_heading(doc, "9. 참조 산출물", 1)
    add_table(
        doc,
        ["산출물", "경로"],
        [
            ["통합 Vivado project", r"C:\Users\YangGeon\SNN_ECG_RESTORE_MODEL_S\vivado_project\SNN_ECG_ModelS_Unified\SNN_ECG_ModelS_Unified.xpr"],
            ["Board bitstream", r"C:\Users\YangGeon\SNN_ECG_RESTORE_MODEL_S\bitstreams\nexys_a7_model_s_smoke_top.bit"],
            ["Core synthesis report", r"C:\Users\YangGeon\SNN_ECG_RESTORE_MODEL_S\reports\synth\model_s_rtl_synth_report.md"],
            ["Board summary report", r"C:\Users\YangGeon\SNN_ECG_RESTORE_MODEL_S\reports\board_smoke\nexys_a7_model_s_smoke_summary.md"],
            ["RTL verification report", r"C:\Users\YangGeon\SNN_ECG_RESTORE_MODEL_S\reports\model_s_rtl\model_s_rtl_final_report.md"],
        ],
        widths=[2500, 6860],
    )

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("Model S FPGA Verification Report")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(100, 100, 100)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
